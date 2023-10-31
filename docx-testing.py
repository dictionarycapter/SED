import docx,os
#
import datetime
import AllFunctions,hashlib
from PyQt5 import QtCore, QtGui, QtWidgets
import hashlib,json,sys,datetime
from cryptography.hazmat.primitives.asymmetric import rsa, padding
from cryptography.hazmat.primitives import serialization, hashes
from cryptography import x509
from cryptography.x509.oid import NameOID
import datetime
from cryptography.exceptions import InvalidSignature
# doc = docx.Document('222.docx')
# count=0
# for i in doc.paragraphs:
#     count+=1
#     print(count, i.text)
# table=doc.tables[0]
#
# for row in table.rows:
#     string = ''
#     for cell in row.cells:
#         string = cell.text
#         print(string)
#         string=''
# input("Press Enter to continue...")


# #Testing of dictionaries
# dictin={'c':1,'b':'pip'}
# d=[]
# d.append(dictin.values())
# d.append(dictin)
# print(d)

# times= datetime.date.today()
# print(times.year,'\n\n\n')

# hesh=AllFunctions.load_HeshData()
# print(hesh[0])
# DiffentDocs=AllFunctions.load_Different_docks()
# HeshData=AllFunctions.load_HeshData()
# open_text_in_str=''
#










self.tableWidget.setColumnWidth(0, 100)
self.tableWidget.setColumnWidth(2, 313)
self.tableWidget.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
try:
    for row in range(self.tableWidget.rowCount()):
        for column in range(self.tableWidget.columnCount()):
            item = self.tableWidget.item(row, column)
            item.setFlags(item.flags() & ~QtCore.Qt.ItemIsEditable)
            if item is not None:
                item.setFlags(item.flags() ^ QtCore.Qt.ItemIsEnabled)
except:
    pass



'''Заполнение данных таблицами'''
striing = 0
rolls = 0
for i in self.DataToTable:
    for j in i:
        item1 = QtWidgets.QTableWidgetItem(str(j))
        self.tableWidget.setItem(striing, rolls, item1)
        rolls += 1
    rolls = 0
    striing += 1



# open_text_in_str='aaaaa'
##
# signature = key.sign(open_text_in_str.encode(), padding.PSS(mgf=padding.MGF1(hashes.SHA256()),
#                                                                   salt_length=padding.PSS.MAX_LENGTH),
#                            hashes.SHA256())
# # '''Изменить название подписи'''
# with open(f'signautres/{user}_signature{NDoc}.bin', 'wb') as f:
#     f.write(signature)
#
#
# cert.public_key()
# try:
#     cert.public_key().verify(
#         signature,
#         open_text_in_str.encode(),
#         padding.PKCS1v15(),
#         hashes.SHA256()
#     )
#     print('All is ok')
# except InvalidSignature:
#     print('Invalid_signature')
#








#
#
#
#
#
# admin_key, admin_cert, user = AllFunctions.Load_Admin_key_and_cert()
# AllFunctions.sign_doc(user,open_text_in_str,admin_key,1)
# user_key, user_cert, user = AllFunctions.Load_User_key_and_cert()
# AllFunctions.sign_doc(user,open_text_in_str,user_key,1)
# print('data_is_ready')
#
# admin_signature = AllFunctions.load_Admin_signature(1)
# user_signature = AllFunctions.load_User_signature(1)
#
# AllFunctions.verify_signature(user_key, open_text_in_str, user_signature)
# AllFunctions.verify_signature(admin_key, open_text_in_str, admin_signature)
# print('ready')